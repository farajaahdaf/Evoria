from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from copy import deepcopy
import os

src  = '/root/.claude/uploads/80c9ac54-b8d0-43e8-baac-9cb9479a5634/67cb85a5-template_portofolio_eo_evoria_1.docx'
dest = '/home/user/Evoria/public/templates/test_filled_portofolio_eo.docx'

doc = Document(src)

# ── helper: ganti isi cell dengan teks nyata ──────────────────────────────────
def fill_cell(cell, text, bold=False):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.bold       = bold
    run.font.italic = False
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)   # hitam gelap, bukan abu hint

# ── Data test realistis ───────────────────────────────────────────────────────
BAG1 = [
    ('Nama Perusahaan / Komunitas *',         'PT Suara Nusantara Event'),
    ('Jenis Organisasi *',                    'PT (Perseroan Terbatas)'),
    ('Tahun Berdiri *',                       '2019'),
    ('NPWP',                                  '12.345.678.9-021.000'),
    ('Alamat Lengkap *',                      'Jl. Jend. Sudirman No. 45, RT 003/RW 005, Kel. Senayan, Kec. Kebayoran Baru'),
    ('Kota / Kabupaten *',                    'Jakarta Selatan'),
    ('Provinsi *',                            'DKI Jakarta'),
    ('Website',                               'https://www.suaranusantara.id'),
    ('Instagram',                             '@suaranusantaraofficial'),
    ('LinkedIn / Facebook',                   'https://linkedin.com/company/suara-nusantara-event'),
]

BAG2 = [
    ('Nama Lengkap PIC *',   'Budi Santoso'),
    ('Jabatan / Posisi *',   'Direktur Utama'),
    ('Nomor HP / WA *',      '081234567890'),
    ('Email *',              'budi@suaranusantara.id'),
    ('Nomor KTP',            '3171234567890001'),
]

DESKRIPSI = (
    'PT Suara Nusantara Event adalah perusahaan event organizer yang berdiri sejak tahun 2019 '
    'dan berbasis di Jakarta Selatan. Kami bergerak di bidang penyelenggaraan event hiburan, '
    'festival musik, seminar nasional, dan pameran seni budaya. Dengan tim yang berpengalaman '
    'lebih dari 7 tahun di industri event, kami telah berhasil menyelenggarakan lebih dari '
    '25 event berskala besar di berbagai kota di Indonesia, mulai dari Jakarta, Bandung, '
    'Surabaya, Yogyakarta, hingga Bali. Kami mengutamakan profesionalisme, keamanan peserta, '
    'dan kualitas pengalaman yang tak terlupakan. Setiap event yang kami selenggarakan dirancang '
    'dengan detail, mulai dari perencanaan venue, manajemen tiket, hingga pengelolaan artis '
    'dan sponsor. Kami percaya bahwa setiap event adalah kesempatan untuk mempertemukan '
    'komunitas dan menciptakan momen bermakna bagi seluruh peserta.'
)

VISI  = 'Menjadi event organizer terpercaya dan terdepan di Indonesia yang menghadirkan pengalaman event berkualitas internasional.'
MISI  = 'Menyelenggarakan event yang aman, profesional, dan berkesan; memberdayakan talenta lokal; serta membangun ekosistem industri hiburan Indonesia yang sehat.'

BAG4_INFO = [
    ('Lama Pengalaman di Bidang Event *',          '7 tahun'),
    ('Total Event yang Pernah Diselenggarakan *',  '25 event'),
    ('Jenis Event yang Biasa Digeluti *',          'Konser Musik, Festival, Seminar Nasional, Pameran Seni'),
    ('Kapasitas Peserta Terbesar',                 '10.000 orang'),
    ('Kota-kota yang Pernah Dijangkau',            'Jakarta, Bandung, Surabaya, Yogyakarta, Bali, Medan'),
]

REKAM_JEJAK = [
    ('Jakarta Music Festival 2024',  '2024', 'Jakarta',    '8.500'),
    ('Bandung Jazz Night 2023',       '2023', 'Bandung',    '3.200'),
    ('Tech Summit Indonesia 2023',    '2023', 'Jakarta',    '2.800'),
    ('Bali Cultural Festival 2022',   '2022', 'Bali',       '5.000'),
    ('Yogya Art Week 2021',           '2021', 'Yogyakarta', '4.100'),
]

BAG5_INFO = [
    ('Judul Event *',                   'Evoria Music Festival 2026'),
    ('Kategori Event *',                'Konser / Festival Musik'),
    ('Perkiraan Tanggal Pelaksanaan *', '20 September 2026'),
    ('Perkiraan Lokasi / Venue *',      'Gelora Bung Karno (GBK), Jakarta Pusat'),
    ('Perkiraan Jumlah Peserta *',      '15.000 orang'),
    ('Kisaran Harga Tiket',             'Rp250.000 – Rp1.500.000'),
]

DESKRIPSI_EVENT = (
    'Evoria Music Festival 2026 adalah festival musik tahunan yang menghadirkan deretan artis '
    'ternama Indonesia dan mancanegara dalam satu panggung megah di GBK Jakarta. '
    'Event ini dirancang untuk 15.000 penonton dengan 3 stage berbeda, area kuliner, '
    'zona merchandise, dan pengalaman interaktif digital. Pembelian tiket dilakukan '
    'secara eksklusif melalui platform Evoria.'
)

TTD = [
    ('Nama Lengkap', 'Budi Santoso'),
    ('Jabatan',      'Direktur Utama'),
    ('Tanggal',      '29 Mei 2026'),
    ('Tanda Tangan', '(tanda tangan terlampir)'),
]

# ── Isi tabel ─────────────────────────────────────────────────────────────────
tables = doc.tables

# TABLE 2 → Bagian 1
for i, (lbl, val) in enumerate(BAG1):
    fill_cell(tables[2].rows[i].cells[1], val)

# TABLE 3 → Bagian 2
for i, (lbl, val) in enumerate(BAG2):
    fill_cell(tables[3].rows[i].cells[1], val)

# TABLE 4 → Deskripsi (single cell, ada \n kosong di dalamnya)
fill_cell(tables[4].rows[0].cells[0], DESKRIPSI)

# TABLE 5 → Visi & Misi
fill_cell(tables[5].rows[0].cells[1], VISI)
fill_cell(tables[5].rows[1].cells[1], MISI)

# TABLE 6 → Bagian 4 info
for i, (lbl, val) in enumerate(BAG4_INFO):
    fill_cell(tables[6].rows[i].cells[1], val)

# TABLE 7 → Rekam Jejak (baris 0 = header, 1–5 = data)
for i, (nama, tahun, kota, peserta) in enumerate(REKAM_JEJAK):
    row = tables[7].rows[i + 1]
    fill_cell(row.cells[1], nama)
    fill_cell(row.cells[2], tahun)
    fill_cell(row.cells[3], kota)
    fill_cell(row.cells[4], peserta)

# TABLE 8 → Bagian 5 info
for i, (lbl, val) in enumerate(BAG5_INFO):
    fill_cell(tables[8].rows[i].cells[1], val)

# TABLE 9 → Deskripsi event (single cell)
fill_cell(tables[9].rows[0].cells[0], DESKRIPSI_EVENT)

# TABLE 11 → TTD
for i, (lbl, val) in enumerate(TTD):
    fill_cell(tables[11].rows[i].cells[1], val)

os.makedirs(os.path.dirname(dest), exist_ok=True)
doc.save(dest)
print(f'Saved → {dest}')
