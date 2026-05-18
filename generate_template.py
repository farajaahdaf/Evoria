from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **borders):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, conf in borders.items():
        tag = OxmlElement(f'w:{edge}')
        for attr, val in conf.items():
            tag.set(qn(f'w:{attr}'), val)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def heading_paragraph(doc, text, level=1):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12) if level == 1 else Pt(10.5)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF) if level == 1 else RGBColor(0x1E, 0x3A, 0x8A)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def field_row(table, label, required=True, hint=''):
    row   = table.add_row()
    badge = ' *' if required else ''
    lbl   = row.cells[0]
    val   = row.cells[1]

    lbl_run = lbl.paragraphs[0].add_run(label + badge)
    lbl_run.font.size = Pt(10)
    lbl_run.bold      = required
    if required:
        lbl_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    hint_run = val.paragraphs[0].add_run(hint if hint else '')
    hint_run.font.size  = Pt(9)
    hint_run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
    hint_run.font.italic    = True

    set_cell_bg(lbl, 'EFF6FF')
    for edge in ('top', 'bottom', 'left', 'right'):
        set_cell_border(lbl, **{edge: {'val': 'single', 'sz': '4', 'color': 'BFDBFE'}})
        set_cell_border(val, **{edge: {'val': 'single', 'sz': '4', 'color': 'BFDBFE'}})

    return row

def two_col_table(doc, widths=(5, 11)):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Cm(widths[0])
    tbl.columns[1].width = Cm(widths[1])
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
# HEADER BANNER
# ══════════════════════════════════════════════════════════════════════════════
header_tbl = doc.add_table(rows=1, cols=1)
header_tbl.style = 'Table Grid'
hcell = header_tbl.rows[0].cells[0]
set_cell_bg(hcell, '1E40AF')
for edge in ('top', 'bottom', 'left', 'right'):
    set_cell_border(hcell, **{edge: {'val': 'none', 'sz': '0', 'color': 'FFFFFF'}})

p1 = hcell.paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p1.add_run('EVORIA')
r1.bold = True
r1.font.size = Pt(28)
r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

p2 = hcell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('TEMPLATE PORTOFOLIO EVENT ORGANIZER')
r2.bold = True
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFE)

p3 = hcell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Versi 1.0  •  Harap diisi lengkap sebelum diunggah ke platform')
r3.font.size = Pt(9)
r3.font.color.rgb = RGBColor(0x93, 0xC5, 0xFD)
p3.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PETUNJUK PENGISIAN
# ══════════════════════════════════════════════════════════════════════════════
notice_tbl = doc.add_table(rows=1, cols=1)
notice_tbl.style = 'Table Grid'
ncell = notice_tbl.rows[0].cells[0]
set_cell_bg(ncell, 'FEF9C3')
for edge in ('top', 'bottom', 'left', 'right'):
    set_cell_border(ncell, **{edge: {'val': 'single', 'sz': '6', 'color': 'FDE047'}})

np0 = ncell.paragraphs[0]
nr0 = np0.add_run('📋  PETUNJUK PENGISIAN')
nr0.bold = True
nr0.font.size = Pt(10)
nr0.font.color.rgb = RGBColor(0x92, 0x40, 0x00)

instructions = [
    'Kolom bertanda  *  adalah WAJIB diisi. Pengajuan tidak akan diproses jika kosong.',
    'Kolom tanpa tanda  *  bersifat opsional, namun meningkatkan skor kelengkapan Anda.',
    'Skor kelengkapan minimum yang direkomendasikan: 80/100.',
    'Hapus teks petunjuk (cetak miring abu-abu) sebelum mengisi kolom.',
    'Simpan file dengan format .docx lalu unggah melalui halaman pengajuan Evoria.',
]
for instr in instructions:
    np = ncell.add_paragraph(style='List Bullet')
    nr = np.add_run(instr)
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x78, 0x35, 0x00)
    np.paragraph_format.space_after = Pt(1)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# BAGIAN 1 – IDENTITAS ORGANISASI
# ══════════════════════════════════════════════════════════════════════════════
heading_paragraph(doc, '📌  BAGIAN 1 — IDENTITAS ORGANISASI')

tbl1 = two_col_table(doc)
field_row(tbl1, 'Nama Perusahaan / Komunitas',  required=True,  hint='Contoh: PT Maju Bersama / Komunitas Jazz Jakarta')
field_row(tbl1, 'Jenis Organisasi',              required=True,  hint='PT / CV / Yayasan / Komunitas / Perorangan')
field_row(tbl1, 'Tahun Berdiri',                 required=True,  hint='Contoh: 2018')
field_row(tbl1, 'NPWP',                          required=False, hint='Contoh: 12.345.678.9-012.000')
field_row(tbl1, 'Alamat Lengkap',                required=True,  hint='Jalan, nomor, RT/RW, kelurahan, kecamatan')
field_row(tbl1, 'Kota / Kabupaten',              required=True,  hint='Contoh: Jakarta Selatan')
field_row(tbl1, 'Provinsi',                      required=True,  hint='Contoh: DKI Jakarta')
field_row(tbl1, 'Website',                       required=False, hint='https://www.namaorganisasi.com')
field_row(tbl1, 'Instagram',                     required=False, hint='@namainstagram')
field_row(tbl1, 'LinkedIn / Facebook',           required=False, hint='URL profil media sosial lainnya')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# BAGIAN 2 – PENANGGUNG JAWAB
# ══════════════════════════════════════════════════════════════════════════════
heading_paragraph(doc, '👤  BAGIAN 2 — PENANGGUNG JAWAB (PIC)')

tbl2 = two_col_table(doc)
field_row(tbl2, 'Nama Lengkap PIC',   required=True,  hint='Sesuai KTP')
field_row(tbl2, 'Jabatan / Posisi',   required=True,  hint='Contoh: Direktur / Ketua Komunitas / Event Manager')
field_row(tbl2, 'Nomor HP / WA',      required=True,  hint='Contoh: 081234567890 (aktif & dapat dihubungi)')
field_row(tbl2, 'Email',              required=True,  hint='Email aktif yang sama dengan akun Evoria')
field_row(tbl2, 'Nomor KTP',          required=False, hint='16 digit NIK sesuai KTP')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# BAGIAN 3 – PROFIL ORGANISASI
# ══════════════════════════════════════════════════════════════════════════════
heading_paragraph(doc, '🏢  BAGIAN 3 — PROFIL ORGANISASI')

p_sub = doc.add_paragraph()
p_sub.add_run('Deskripsi Organisasi  *').bold = True
p_sub.runs[0].font.size = Pt(10)
p_sub.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
p_sub.paragraph_format.space_after = Pt(2)

p_hint = doc.add_paragraph()
r_hint = p_hint.add_run('Jelaskan secara singkat tentang organisasi Anda: bidang yang digeluti, nilai utama, keunggulan, dan target peserta event. Minimal 100 kata.')
r_hint.font.size = Pt(9)
r_hint.font.italic = True
r_hint.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
p_hint.paragraph_format.space_after = Pt(2)

desc_tbl = doc.add_table(rows=1, cols=1)
desc_tbl.style = 'Table Grid'
dcell = desc_tbl.rows[0].cells[0]
set_cell_bg(dcell, 'F8FAFC')
for edge in ('top', 'bottom', 'left', 'right'):
    set_cell_border(dcell, **{edge: {'val': 'single', 'sz': '4', 'color': 'BFDBFE'}})
dp = dcell.paragraphs[0]
dr = dp.add_run('\n\n\n\n\n')
dr.font.size = Pt(10)

doc.add_paragraph()

p_sub2 = doc.add_paragraph()
p_sub2.add_run('Visi & Misi').bold = True
p_sub2.runs[0].font.size = Pt(10)
p_sub2.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
p_sub2.paragraph_format.space_after = Pt(2)

vm_tbl = doc.add_table(rows=2, cols=2)
vm_tbl.style = 'Table Grid'
for i, (lbl, hint) in enumerate([('Visi', 'Contoh: Menjadi platform event terpercaya...'), ('Misi', 'Contoh: Menghadirkan pengalaman event berkualitas...')]):
    set_cell_bg(vm_tbl.rows[i].cells[0], 'EFF6FF')
    r = vm_tbl.rows[i].cells[0].paragraphs[0].add_run(lbl)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    rh = vm_tbl.rows[i].cells[1].paragraphs[0].add_run(hint)
    rh.font.size = Pt(9)
    rh.font.italic = True
    rh.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
    for edge in ('top', 'bottom', 'left', 'right'):
        set_cell_border(vm_tbl.rows[i].cells[0], **{edge: {'val': 'single', 'sz': '4', 'color': 'BFDBFE'}})
        set_cell_border(vm_tbl.rows[i].cells[1], **{edge: {'val': 'single', 'sz': '4', 'color': 'BFDBFE'}})

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# BAGIAN 4 – PENGALAMAN
# ══════════════════════════════════════════════════════════════════════════════
heading_paragraph(doc, '🏆  BAGIAN 4 — PENGALAMAN MENYELENGGARAKAN EVENT')

tbl3 = two_col_table(doc)
field_row(tbl3, 'Lama Pengalaman di Bidang Event', required=True,  hint='Contoh: 5 tahun')
field_row(tbl3, 'Total Event yang Pernah Diselenggarakan', required=True, hint='Contoh: 12 event')
field_row(tbl3, 'Jenis Event yang Biasa Digeluti', required=True,
          hint='Konser / Seminar / Festival / Olahraga / Pameran / Lainnya')
field_row(tbl3, 'Kapasitas Peserta Terbesar',       required=False, hint='Contoh: 5.000 orang')
field_row(tbl3, 'Kota-kota yang Pernah Dijangkau',  required=False, hint='Contoh: Jakarta, Bandung, Surabaya')

doc.add_paragraph()

p_track = doc.add_paragraph()
r_track = p_track.add_run('Rekam Jejak Event (isi minimal 3 event)  *')
r_track.bold = True
r_track.font.size = Pt(10)
r_track.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
p_track.paragraph_format.space_after = Pt(4)

track_tbl = doc.add_table(rows=6, cols=5)
track_tbl.style = 'Table Grid'
headers = ['No', 'Nama Event', 'Tahun', 'Kota', 'Jml Peserta']
widths  = [1, 6, 2, 3, 3]
for i, (h, w) in enumerate(zip(headers, widths)):
    track_tbl.columns[i].width = Cm(w)
    cell = track_tbl.rows[0].cells[i]
    set_cell_bg(cell, '1E40AF')
    for edge in ('top', 'bottom', 'left', 'right'):
        set_cell_border(cell, **{edge: {'val': 'single', 'sz': '4', 'color': '1E3A8A'}})
    rh = cell.paragraphs[0].add_run(h)
    rh.bold = True
    rh.font.size = Pt(9)
    rh.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for row_idx in range(1, 6):
    for col_idx in range(5):
        cell = track_tbl.rows[row_idx].cells[col_idx]
        bg   = 'F8FAFC' if row_idx % 2 == 0 else 'FFFFFF'
        set_cell_bg(cell, bg)
        for edge in ('top', 'bottom', 'left', 'right'):
            set_cell_border(cell, **{edge: {'val': 'single', 'sz': '4', 'color': 'BFDBFE'}})
        if col_idx == 0:
            cell.paragraphs[0].add_run(str(row_idx)).font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# BAGIAN 5 – RENCANA EVENT
# ══════════════════════════════════════════════════════════════════════════════
heading_paragraph(doc, '📅  BAGIAN 5 — RENCANA EVENT DI EVORIA')

p_note = doc.add_paragraph()
r_note = p_note.add_run('Jelaskan event pertama yang akan Anda buat di platform Evoria setelah pengajuan disetujui.')
r_note.font.size  = Pt(9)
r_note.font.italic = True
r_note.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
p_note.paragraph_format.space_after = Pt(4)

tbl4 = two_col_table(doc)
field_row(tbl4, 'Judul Event',                    required=True,  hint='Contoh: Jakarta Jazz Night 2026')
field_row(tbl4, 'Kategori Event',                 required=True,  hint='Konser / Seminar / Festival / Olahraga / dll')
field_row(tbl4, 'Perkiraan Tanggal Pelaksanaan',  required=True,  hint='Contoh: 15 Agustus 2026')
field_row(tbl4, 'Perkiraan Lokasi / Venue',       required=True,  hint='Contoh: Istora Senayan, Jakarta')
field_row(tbl4, 'Perkiraan Jumlah Peserta',       required=True,  hint='Contoh: 2.000 orang')
field_row(tbl4, 'Kisaran Harga Tiket',            required=False, hint='Contoh: Rp150.000 – Rp500.000 / Gratis')

doc.add_paragraph()

p_desc_ev = doc.add_paragraph()
p_desc_ev.add_run('Deskripsi Singkat Event  *').bold = True
p_desc_ev.runs[0].font.size = Pt(10)
p_desc_ev.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
p_desc_ev.paragraph_format.space_after = Pt(2)

ev_tbl = doc.add_table(rows=1, cols=1)
ev_tbl.style = 'Table Grid'
evcell = ev_tbl.rows[0].cells[0]
set_cell_bg(evcell, 'F8FAFC')
for edge in ('top', 'bottom', 'left', 'right'):
    set_cell_border(evcell, **{edge: {'val': 'single', 'sz': '4', 'color': 'BFDBFE'}})
evp = evcell.paragraphs[0]
evr = evp.add_run('\n\n\n\n')
evr.font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# BAGIAN 6 – PERNYATAAN & TTD
# ══════════════════════════════════════════════════════════════════════════════
heading_paragraph(doc, '✍️  BAGIAN 6 — PERNYATAAN DAN TANDA TANGAN')

decl_tbl = doc.add_table(rows=1, cols=1)
decl_tbl.style = 'Table Grid'
dcell2 = decl_tbl.rows[0].cells[0]
set_cell_bg(dcell2, 'EFF6FF')
for edge in ('top', 'bottom', 'left', 'right'):
    set_cell_border(dcell2, **{edge: {'val': 'single', 'sz': '6', 'color': '1E40AF'}})

dp1 = dcell2.paragraphs[0]
dr1 = dp1.add_run(
    'Saya yang bertanda tangan di bawah ini menyatakan bahwa seluruh informasi yang tercantum '
    'dalam dokumen ini adalah benar, lengkap, dan dapat dipertanggungjawabkan. Saya bersedia '
    'menanggung segala konsekuensi hukum apabila terdapat informasi yang tidak sesuai dengan '
    'kenyataan.'
)
dr1.font.size = Pt(10)
dr1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

doc.add_paragraph()

sign_tbl = doc.add_table(rows=4, cols=2)
sign_tbl.style = 'Table Grid'
sign_labels = ['Nama Lengkap', 'Jabatan', 'Tanggal', 'Tanda Tangan']
for i, lbl in enumerate(sign_labels):
    lc = sign_tbl.rows[i].cells[0]
    rc = sign_tbl.rows[i].cells[1]
    set_cell_bg(lc, 'EFF6FF')
    for edge in ('top', 'bottom', 'left', 'right'):
        set_cell_border(lc, **{edge: {'val': 'single', 'sz': '4', 'color': 'BFDBFE'}})
        set_cell_border(rc, **{edge: {'val': 'single', 'sz': '4', 'color': 'BFDBFE'}})
    r = lc.paragraphs[0].add_run(lbl)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    if lbl == 'Tanda Tangan':
        rc_r = rc.paragraphs[0].add_run('\n\n\n')
        rc_r.font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SCORING REFERENCE (info box)
# ══════════════════════════════════════════════════════════════════════════════
heading_paragraph(doc, '📊  REFERENSI SKOR KELENGKAPAN', level=2)

score_tbl = doc.add_table(rows=5, cols=3)
score_tbl.style = 'Table Grid'

score_headers = ['Bagian', 'Bobot', 'Keterangan']
score_data = [
    ('Identitas Organisasi (Bag. 1)',   '25 poin', 'Nama, jenis, alamat, kota, provinsi wajib diisi'),
    ('Penanggung Jawab (Bag. 2)',        '20 poin', 'Nama PIC, jabatan, HP, email wajib diisi'),
    ('Profil Organisasi (Bag. 3)',       '20 poin', 'Deskripsi min. 100 kata'),
    ('Pengalaman + Rekam Jejak (Bag. 4)','20 poin', 'Min. 3 event di tabel rekam jejak'),
    ('Rencana Event (Bag. 5)',           '15 poin', 'Semua field wajib di bagian ini terisi'),
]

for i, (b, w, k) in enumerate(score_data):
    row = score_tbl.rows[i]  # no header row, data starts at 0
    bg = 'F0F9FF' if i % 2 == 0 else 'FFFFFF'
    for ci, txt in enumerate([b, w, k]):
        cell = row.cells[ci]
        set_cell_bg(cell, bg)
        for edge in ('top', 'bottom', 'left', 'right'):
            set_cell_border(cell, **{edge: {'val': 'single', 'sz': '4', 'color': 'BAE6FD'}})
        r = cell.paragraphs[0].add_run(txt)
        r.font.size = Pt(9)
        if ci == 1:
            r.bold = True
            r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER NOTE
# ══════════════════════════════════════════════════════════════════════════════
footer_tbl = doc.add_table(rows=1, cols=1)
footer_tbl.style = 'Table Grid'
fcell = footer_tbl.rows[0].cells[0]
set_cell_bg(fcell, '1E3A8A')
for edge in ('top', 'bottom', 'left', 'right'):
    set_cell_border(fcell, **{edge: {'val': 'none', 'sz': '0', 'color': 'FFFFFF'}})
fp = fcell.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run('© 2026 Evoria  •  Template ini hanya berlaku untuk pengajuan melalui platform Evoria  •  evoria.id')
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x93, 0xC5, 0xFD)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = '/home/user/Evoria/public/templates/template_portofolio_eo_evoria.docx'
import os
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'Template saved → {output_path}')
